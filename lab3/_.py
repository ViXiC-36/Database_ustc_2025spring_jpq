import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
from db import DatabaseConnection
import datetime

class TeacherResearchSystem:
    def __init__(self, root):
        self.root = root
        self.root.title("教师教学科研登记系统")
        self.root.geometry("1000x700")
        self.db = DatabaseConnection()

        # ! 美化
        self.fonts = {
            "default": ("华文仿宋", 12),
            "title": ("汉仪风骨楷体 W", 30),
            "subtitle": ("汉仪风骨楷体 W", 16),
            "button": ("华文仿宋", 12),
            "label": ("华文仿宋", 12),
        }
        self.colors = {
            "bg": "#fff6e5",
            "dk": "#3e454c",
            "db": "#2185c5",
            "lb": "#7ecefd",
            "rd": "#ff7f66",
        }
        
        self.create_menu()
        self.create_main_interface()
    
    def create_menu(self):
        # menu
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # 教师管理菜单
        teacher_menu = tk.Menu(menubar, tearoff=0)
        # ? add_cascade 方法用于在菜单栏中添加一个下拉菜单
        menubar.add_cascade(label="教师管理", menu=teacher_menu)
        teacher_menu.add_command(label="新教师", command=self.add_teacher)
        # teacher_menu.add_command(label="查询教师", command=self.search_teacher)
        
        # 论文管理菜单
        paper_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="论文管理", menu=paper_menu)
        paper_menu.add_command(label="新论文", command=self.add_paper)
        paper_menu.add_command(label="论文发表登记", command=self.register_paper)
        paper_menu.add_command(label="查询论文", command=self.search_paper)
        
        # 项目管理菜单
        project_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="项目管理", menu=project_menu)
        project_menu.add_command(label="添加项目", command=self.add_project)
        project_menu.add_command(label="项目承担登记", command=self.register_project)
        project_menu.add_command(label="查询项目", command=self.search_project)
        
        # 课程管理菜单
        course_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="课程管理", menu=course_menu)
        course_menu.add_command(label="添加课程", command=self.add_course)
        course_menu.add_command(label="授课登记", command=self.register_teaching)
        course_menu.add_command(label="查询课程", command=self.search_course)
        
        # 统计查询菜单
        stats_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="统计查询", menu=stats_menu)
        stats_menu.add_command(label="年度科研统计", command=self.yearly_stats)
        stats_menu.add_command(label="生成工作量统计表", command=self.generate_workload_report)
    
    def create_main_interface(self):
        """创建主界面"""
        # 欢迎标签
        welcome_label = tk.Label(
            self.root, 
            text="教师教学科研登记系统", 
            font=self.fonts["title"],
            fg=self.colors["db"],
            bg=self.colors["bg"],
        )
        # ? 标签放在中间偏上，上下留出50px的空隙
        # ? pack 是最简单的布局管理器之一，能把组件从上往下依次排列
        welcome_label.pack(pady=50)
        
        # 功能按钮
        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=20)
        
        buttons = [
            ("教师管理", self.search_teacher),
            ("论文管理", self.search_paper), 
            ("项目管理", self.search_project),
            ("课程管理", self.search_course),
            ("统计查询", self.yearly_stats)
        ]
        
        for i, (text, command) in enumerate(buttons):
            btn = tk.Button(
                button_frame, 
                text=text, 
                command=command,
                width=15, 
                height=2, 
                font=self.fonts["button"],
                bg=self.colors["lb"],
                fg=self.colors["dk"],
                activebackground=self.colors["db"], # 按钮被点击时的背景色
                activeforeground=self.colors["bg"], # 按钮被点击时的前景色
                relief="groove",    # 凸起的边框
                bd=2,   # 边框宽度
                cursor="hand2"  # 鼠标悬停时的手型光标
            )
            btn.grid(row=i, column=1, padx=10, pady=10)
    
    def add_teacher(self):
        """添加教师"""
        # ? Toplevel 创建一个新的窗口
        # ? title 设置窗口标题
        # ? geometry 设置窗口大小
        # ? configure 设置窗口背景色
        # ? resizable 设置窗口是否可以调整大小
        # ? bg 设置背景色
        add_window = tk.Toplevel(self.root)
        add_window.title("添加教师")
        add_window.geometry("400x300")
        add_window.configure(bg=self.colors["bg"])  
        add_window.resizable(False, False)  # 禁止调整窗口大小

        # 标题标签
        title_label = tk.Label(
            add_window, 
            text="添加教师信息", 
            font=self.fonts["subtitle"],
            fg=self.colors["db"],
            bg=self.colors["bg"],
            anchor='center'  # 居中对齐
        )
        title_label.grid(row=0, column=0, columnspan=2, pady=5)
        
        # 输入字段样式
        label_style = {
            "font": self.fonts["label"],
            "bg": self.colors["bg"],
            "fg": self.colors["dk"]
        }
        
        entry_style = {
            "width": 25,
            "font": self.fonts["default"],
            "relief": "solid",
            "bd": 1,
            "highlightthickness": 2,
            "highlightcolor": self.colors["dk"],
        }
        
        # 输入字段
        # ? Label 用于显示文本，text 参数设置标签内容
        # ? grid 方法用于将组件放置在网格布局中
        # ? sticky='w' 表示左对齐，padx 和 pady 用于设置内边距
        # ? sticky的取值可以是 'n', 's', 'e', 'w'，分别表示北、南、东、西
        # ? **label_style 用于解包字典，将字典中的键值对作为参数传递给 Label
        tk.Label(add_window, text="工号:", **label_style).grid(row=0, column=0, sticky='w', padx=10, pady=5)
        # ? Entry 用于输入文本，width 设置宽度
        # ? grid 方法用于将组件放置在网格布局中
        teacher_id_entry = tk.Entry(add_window, width=30)
        teacher_id_entry.grid(row=0, column=1, padx=10, pady=5)
        
        tk.Label(add_window, text="姓名:", **label_style).grid(row=1, column=0, sticky='w', padx=10, pady=5)
        name_entry = tk.Entry(add_window, width=30)
        name_entry.grid(row=1, column=1, padx=10, pady=5)
        
        tk.Label(add_window, text="性别:", **label_style).grid(row=2, column=0, sticky='w', padx=10, pady=5)
        gender_var = tk.StringVar()
        gender_combo = ttk.Combobox(add_window, textvariable=gender_var, width=27,
                                   values=["1-男", "2-女"], state='readonly')
        gender_combo.grid(row=2, column=1, padx=10, pady=5)
        
        tk.Label(add_window, text="职称:", **label_style).grid(row=3, column=0, sticky='w', padx=10, pady=5)
        title_var = tk.StringVar()
        title_combo = ttk.Combobox(add_window, textvariable=title_var, width=27,
                                  values=["1-博士后", "2-助教", "3-讲师", "4-副教授", 
                                         "5-特任教授", "6-教授", "7-助理研究员", 
                                         "8-特任副研究员", "9-副研究员", "10-特任研究员", "11-研究员"],
                                  state='readonly')
        title_combo.grid(row=3, column=1, padx=10, pady=5)
        
        def save_teacher():
            teacher_id = teacher_id_entry.get().strip()
            name = name_entry.get().strip()
            gender = gender_var.get()
            title = title_var.get()
            
            # 数据验证
            if not all([teacher_id, name, gender, title]):
                messagebox.showerror("错误", "所有字段都必须填写")
                return
            
            if len(teacher_id) > 5:
                messagebox.showerror("错误", "工号不能超过5个字符")
                return
            
            # 提取数字
            gender_num = int(gender.split('-')[0])
            title_num = int(title.split('-')[0])
            
            # 插入数据库
            query = "INSERT INTO Teacher (teacher_id, name, gender, title) VALUES (%s, %s, %s, %s)"
            result = self.db.execute_update(query, (teacher_id, name, gender_num, title_num))
            
            if result > 0:
                messagebox.showinfo("成功", "教师添加成功")
                add_window.destroy()
            else:
                messagebox.showerror("错误", "教师添加失败，工号可能已存在")
        
        tk.Button(add_window, text="保存", command=save_teacher).grid(row=4, column=0, columnspan=2, pady=20)
    
    def add_paper(self):
        """添加论文"""
        add_window = tk.Toplevel(self.root)
        add_window.title("添加论文")
        add_window.geometry("500x400")
        
        tk.Label(add_window, text="论文序号:").grid(row=0, column=0, sticky='w', padx=10, pady=5)
        paper_id_entry = tk.Entry(add_window, width=30)
        paper_id_entry.grid(row=0, column=1, padx=10, pady=5)
        
        tk.Label(add_window, text="论文名称:").grid(row=1, column=0, sticky='w', padx=10, pady=5)
        paper_name_entry = tk.Entry(add_window, width=30)
        paper_name_entry.grid(row=1, column=1, padx=10, pady=5)
        
        tk.Label(add_window, text="发表源:").grid(row=2, column=0, sticky='w', padx=10, pady=5)
        pub_source_entry = tk.Entry(add_window, width=30)
        pub_source_entry.grid(row=2, column=1, padx=10, pady=5)
        
        tk.Label(add_window, text="发表年份:").grid(row=3, column=0, sticky='w', padx=10, pady=5)
        pub_year_entry = tk.Entry(add_window, width=30)
        pub_year_entry.grid(row=3, column=1, padx=10, pady=5)
        
        tk.Label(add_window, text="论文类型:").grid(row=4, column=0, sticky='w', padx=10, pady=5)
        paper_type_var = tk.StringVar()
        paper_type_combo = ttk.Combobox(add_window, textvariable=paper_type_var, width=27,
                                       values=["1-full paper", "2-short paper", "3-poster paper", "4-demo paper"],
                                       state='readonly')
        paper_type_combo.grid(row=4, column=1, padx=10, pady=5)
        
        tk.Label(add_window, text="论文级别:").grid(row=5, column=0, sticky='w', padx=10, pady=5)
        paper_level_var = tk.StringVar()
        paper_level_combo = ttk.Combobox(add_window, textvariable=paper_level_var, width=27,
                                        values=["1-CCF-A", "2-CCF-B", "3-CCF-C", "4-中文CCF-A", "5-中文CCF-B", "6-无级别"],
                                        state='readonly')
        paper_level_combo.grid(row=5, column=1, padx=10, pady=5)
        
        def save_paper():
            # 获取数据
            paper_id = paper_id_entry.get().strip()
            paper_name = paper_name_entry.get().strip()
            pub_source = pub_source_entry.get().strip()
            pub_year = pub_year_entry.get().strip()
            paper_type = paper_type_var.get()
            paper_level = paper_level_var.get()
            
            # 数据验证
            if not all([paper_id, paper_name, pub_source, pub_year, paper_type, paper_level]):
                messagebox.showerror("错误", "所有字段都必须填写")
                return
            
            try:
                paper_id = int(paper_id)
                pub_year = int(pub_year)
                paper_type_num = int(paper_type.split('-')[0])
                paper_level_num = int(paper_level.split('-')[0])
            except ValueError:
                messagebox.showerror("错误", "序号和年份必须是数字")
                return
            
            # 插入数据库
            query = """INSERT INTO Paper (paper_id, paper_name, pub_source, pub_year, paper_type, paper_level) 
                      VALUES (%s, %s, %s, %s, %s, %s)"""
            result = self.db.execute_update(query, (paper_id, paper_name, pub_source, pub_year, paper_type_num, paper_level_num))
            
            if result > 0:
                messagebox.showinfo("成功", "论文添加成功")
                add_window.destroy()
            else:
                messagebox.showerror("错误", "论文添加失败")
        
        tk.Button(add_window, text="保存", command=save_paper).grid(row=6, column=0, columnspan=2, pady=20)
    
    def register_paper(self):
        """论文发表登记"""
        register_window = tk.Toplevel(self.root)
        register_window.title("论文发表登记")
        register_window.geometry("400x300")
        
        tk.Label(register_window, text="教师工号:").grid(row=0, column=0, sticky='w', padx=10, pady=5)
        teacher_id_entry = tk.Entry(register_window, width=30)
        teacher_id_entry.grid(row=0, column=1, padx=10, pady=5)
        
        tk.Label(register_window, text="论文序号:").grid(row=1, column=0, sticky='w', padx=10, pady=5)
        paper_id_entry = tk.Entry(register_window, width=30)
        paper_id_entry.grid(row=1, column=1, padx=10, pady=5)
        
        tk.Label(register_window, text="作者排名:").grid(row=2, column=0, sticky='w', padx=10, pady=5)
        author_rank_entry = tk.Entry(register_window, width=30)
        author_rank_entry.grid(row=2, column=1, padx=10, pady=5)
        
        tk.Label(register_window, text="是否通讯作者:").grid(row=3, column=0, sticky='w', padx=10, pady=5)
        is_corresponding_var = tk.BooleanVar()
        is_corresponding_check = tk.Checkbox(register_window, variable=is_corresponding_var)
        is_corresponding_check.grid(row=3, column=1, sticky='w', padx=10, pady=5)
        
        def save_register():
            teacher_id = teacher_id_entry.get().strip()
            paper_id = paper_id_entry.get().strip()
            author_rank = author_rank_entry.get().strip()
            is_corresponding = is_corresponding_var.get()
            
            if not all([teacher_id, paper_id, author_rank]):
                messagebox.showerror("错误", "所有字段都必须填写")
                return
            
            try:
                paper_id = int(paper_id)
                author_rank = int(author_rank)
            except ValueError:
                messagebox.showerror("错误", "论文序号和排名必须是数字")
                return
            
            # 检查教师和论文是否存在
            teacher_check = self.db.execute_query("SELECT teacher_id FROM Teacher WHERE teacher_id = %s", (teacher_id,))
            paper_check = self.db.execute_query("SELECT paper_id FROM Paper WHERE paper_id = %s", (paper_id,))
            
            if not teacher_check:
                messagebox.showerror("错误", "教师不存在")
                return
            if not paper_check:
                messagebox.showerror("错误", "论文不存在")
                return
            
            # 插入数据库
            query = """INSERT INTO Teacher_Paper (teacher_id, paper_id, author_rank, is_corresponding) 
                      VALUES (%s, %s, %s, %s)"""
            result = self.db.execute_update(query, (teacher_id, paper_id, author_rank, is_corresponding))
            
            if result > 0:
                messagebox.showinfo("成功", "论文发表登记成功")
                register_window.destroy()
            else:
                messagebox.showerror("错误", "论文发表登记失败")
        
        tk.Button(register_window, text="保存", command=save_register).grid(row=4, column=0, columnspan=2, pady=20)
    
    def yearly_stats(self):
        """年度科研统计"""
        stats_window = tk.Toplevel(self.root)
        stats_window.title("年度科研统计")
        stats_window.geometry("800x600")
        
        # 输入框架
        input_frame = tk.Frame(stats_window)
        input_frame.pack(pady=10)
        
        tk.Label(input_frame, text="教师工号:").grid(row=0, column=0, padx=5)
        teacher_id_entry = tk.Entry(input_frame, width=15)
        teacher_id_entry.grid(row=0, column=1, padx=5)
        
        tk.Label(input_frame, text="统计年份:").grid(row=0, column=2, padx=5)
        year_entry = tk.Entry(input_frame, width=15)
        year_entry.grid(row=0, column=3, padx=5)
        
        # 结果显示
        result_text = tk.Text(stats_window, width=90, height=30)
        result_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 滚动条
        scrollbar = tk.Scrollbar(result_text)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        result_text.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=result_text.yview)
        
        def search_stats():
            teacher_id = teacher_id_entry.get().strip()
            year = year_entry.get().strip()
            
            if not teacher_id:
                messagebox.showerror("错误", "请输入教师工号")
                return
            
            result_text.delete(1.0, tk.END)
            
            # 获取教师基本信息
            teacher_query = """
            SELECT name, 
                   CASE gender WHEN 1 THEN '男' WHEN 2 THEN '女' END as gender,
                   CASE title 
                        WHEN 1 THEN '博士后' WHEN 2 THEN '助教' WHEN 3 THEN '讲师' 
                        WHEN 4 THEN '副教授' WHEN 5 THEN '特任教授' WHEN 6 THEN '教授'
                        WHEN 7 THEN '助理研究员' WHEN 8 THEN '特任副研究员' WHEN 9 THEN '副研究员'
                        WHEN 10 THEN '特任研究员' WHEN 11 THEN '研究员'
                   END as title
            FROM Teacher WHERE teacher_id = %s
            """
            teacher_info = self.db.execute_query(teacher_query, (teacher_id,))
            
            if not teacher_info:
                messagebox.showerror("错误", "教师不存在")
                return
            
            name, gender, title = teacher_info[0]
            result_text.insert(tk.END, f"教师教学科研工作统计 ({year if year else '全部年份'})\n")
            result_text.insert(tk.END, "="*60 + "\n\n")
            result_text.insert(tk.END, "教师基本信息\n")
            result_text.insert(tk.END, "-"*30 + "\n")
            result_text.insert(tk.END, f"工号: {teacher_id}    姓名: {name}    性别: {gender}    职称: {title}\n\n")
            
            # 教学情况统计
            result_text.insert(tk.END, "教学情况\n")
            result_text.insert(tk.END, "-"*30 + "\n")
            
            course_query = """
            SELECT c.course_id, c.course_name, tc.my_hours, 
                   CASE tc.semester WHEN 1 THEN '春' WHEN 2 THEN '夏' WHEN 3 THEN '秋' END as semester,
                   tc.year
            FROM Teacher_Course tc
            JOIN Course c ON tc.course_id = c.course_id
            WHERE tc.teacher_id = %s
            """
            params = [teacher_id]
            if year:
                course_query += " AND tc.year = %s"
                params.append(year)
            course_query += " ORDER BY tc.year, tc.semester"
            
            courses = self.db.execute_query(course_query, params)
            if courses:
                for course_id, course_name, my_hours, semester, course_year in courses:
                    result_text.insert(tk.END, f"课程号: {course_id}  课程名: {course_name}  主讲学时: {my_hours}  学期: {course_year} {semester}\n")
            else:
                result_text.insert(tk.END, "无教学记录\n")
            
            result_text.insert(tk.END, "\n")
            
            # 论文发表情况
            result_text.insert(tk.END, "发表论文情况\n")
            result_text.insert(tk.END, "-"*30 + "\n")
            
            paper_query = """
            SELECT p.paper_name, p.pub_source, p.pub_year, tp.author_rank, tp.is_corresponding,
                   CASE p.paper_level 
                        WHEN 1 THEN 'CCF-A' WHEN 2 THEN 'CCF-B' WHEN 3 THEN 'CCF-C'
                        WHEN 4 THEN '中文CCF-A' WHEN 5 THEN '中文CCF-B' WHEN 6 THEN '无级别'
                   END as level
            FROM Teacher_Paper tp
            JOIN Paper p ON tp.paper_id = p.paper_id
            WHERE tp.teacher_id = %s
            """
            params = [teacher_id]
            if year:
                paper_query += " AND p.pub_year = %s"
                params.append(year)
            paper_query += " ORDER BY p.pub_year DESC"
            
            papers = self.db.execute_query(paper_query, params)
            if papers:
                for i, (paper_name, pub_source, pub_year, author_rank, is_corresponding, level) in enumerate(papers, 1):
                    corresponding_text = "通讯作者" if is_corresponding else f"排名第{author_rank}"
                    result_text.insert(tk.END, f"{i}. {paper_name}, {pub_source}, {pub_year}, {level}, {corresponding_text}\n")
            else:
                result_text.insert(tk.END, "无论文发表记录\n")
            
            result_text.insert(tk.END, "\n")
            
            # 承担项目情况
            result_text.insert(tk.END, "承担项目情况\n") 
            result_text.insert(tk.END, "-"*30 + "\n")
            
            project_query = """
            SELECT p.project_name, p.project_source, tp.principal_rank, tp.my_fund, p.total_fund,
                   CASE p.project_type
                        WHEN 1 THEN '国家级项目' WHEN 2 THEN '省部级项目' WHEN 3 THEN '市厅级项目'
                        WHEN 4 THEN '企业合作项目' WHEN 5 THEN '其它类型项目'
                   END as project_type,
                   CONCAT(p.start_year, '-', p.end_year) as duration
            FROM Teacher_Project tp
            JOIN Project p ON tp.project_id = p.project_id
            WHERE tp.teacher_id = %s
            """
            params = [teacher_id]
            if year:
                project_query += " AND %s BETWEEN p.start_year AND p.end_year"
                params.append(year)
            
            projects = self.db.execute_query(project_query, params)
            if projects:
                for i, (project_name, project_source, principal_rank, my_fund, total_fund, project_type, duration) in enumerate(projects, 1):
                    principal_text = "项目负责人" if principal_rank == 1 else f"排名第{principal_rank}"
                    result_text.insert(tk.END, f"{i}. {project_name}, {project_source}, {project_type}, {duration}\n")
                    result_text.insert(tk.END, f"    总经费: {total_fund}, 承担经费: {my_fund}, {principal_text}\n")
            else:
                result_text.insert(tk.END, "无项目承担记录\n")
        
        tk.Button(input_frame, text="查询统计", command=search_stats).grid(row=0, column=4, padx=10)
    
    def generate_workload_report(self):
        """生成工作量统计表"""
        teacher_id = simpledialog.askstring("输入", "请输入教师工号:")
        year_range = simpledialog.askstring("输入", "请输入年份范围(如: 2022-2023):")
        
        if not teacher_id or not year_range:
            return
        
        try:
            start_year, end_year = year_range.split('-')
            start_year, end_year = int(start_year), int(end_year)
        except:
            messagebox.showerror("错误", "年份格式错误，请使用格式: 2022-2023")
            return
        
        # 这里可以使用 python-docx 生成 Word 文档
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading(f'教师教学科研工作统计 ({year_range})', 0)
            
            # 获取教师信息并添加到文档
            teacher_query = """
            SELECT name, 
                   CASE gender WHEN 1 THEN '男' WHEN 2 THEN '女' END as gender,
                   CASE title 
                        WHEN 1 THEN '博士后' WHEN 2 THEN '助教' WHEN 3 THEN '讲师' 
                        WHEN 4 THEN '副教授' WHEN 5 THEN '特任教授' WHEN 6 THEN '教授'
                        WHEN 7 THEN '助理研究员' WHEN 8 THEN '特任副研究员' WHEN 9 THEN '副研究员'
                        WHEN 10 THEN '特任研究员' WHEN 11 THEN '研究员'
                   END as title
            FROM Teacher WHERE teacher_id = %s
            """
            teacher_info = self.db.execute_query(teacher_query, (teacher_id,))
            
            if teacher_info:
                name, gender, title = teacher_info[0]
                
                doc.add_heading('教师基本信息', level=1)
                p = doc.add_paragraph()
                p.add_run(f'工号: {teacher_id}    姓名: {name}    性别: {gender}    职称: {title}')
                
                # 可以继续添加教学、科研等详细信息...
                
                # 保存文档
                filename = f"教师工作量统计_{teacher_id}_{year_range}.docx"
                doc.save(filename)
                messagebox.showinfo("成功", f"报告已生成：{filename}")
            
        except ImportError:
            messagebox.showerror("错误", "请安装 python-docx: pip install python-docx")
    
    def search_teacher(self):
        """查询教师"""
        # 类似的实现...
        pass
    
    def search_paper(self):
        """查询论文"""
        # 类似的实现...
        pass
    
    def add_project(self):
        """添加项目"""
        # 类似于 add_paper 的实现
        pass
    
    def register_project(self):
        """项目承担登记"""
        # 类似于 register_paper 的实现
        pass
    
    def search_project(self):
        """查询项目"""
        pass
    
    def add_course(self):
        """添加课程"""
        # 类似于 add_paper 的实现
        pass
    
    def register_teaching(self):
        """授课登记"""
        # 类似于 register_paper 的实现
        pass
    
    def search_course(self):
        """查询课程"""
        pass

if __name__ == "__main__":
    root = tk.Tk()
    app = TeacherResearchSystem(root)
    root.configure(bg="#fff6e5")  # 设置背景色
    root.title("教师教学科研登记系统")
    root.geometry("1000x700+200+50")
    root.mainloop()
